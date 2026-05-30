"""Multi-format document loaders.

Each loader returns a list of `(text, metadata)` tuples — typically one per
"page" (PDF) or sheet (XLSX) or full document. Metadata extracted varies by
format but every loader emits at minimum: `source_file`, `mime`, `loader`.

Supported:  .pdf  .docx  .xlsx .xls  .txt .md  .html .htm  .json  .eml  url
"""
from __future__ import annotations

import email
import json
import mimetypes
from pathlib import Path
from typing import Any
from urllib.parse import urlparse

import httpx


# ---------- PDF ----------

def load_pdf(path: Path) -> list[tuple[str, dict[str, Any]]]:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    meta = reader.metadata or {}
    base_meta = {
        "source_file": path.name,
        "mime": "application/pdf",
        "loader": "pypdf",
        "pages": len(reader.pages),
        "title": str(meta.get("/Title") or "")[:200] if meta else "",
        "author": str(meta.get("/Author") or "")[:200] if meta else "",
    }
    out: list[tuple[str, dict[str, Any]]] = []
    for i, page in enumerate(reader.pages):
        text = (page.extract_text() or "").strip()
        if not text:
            continue
        out.append((text, {**base_meta, "page_number": i + 1}))
    return out


# ---------- DOCX ----------

def load_docx(path: Path) -> list[tuple[str, dict[str, Any]]]:
    from docx import Document

    doc = Document(str(path))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    # tables too — they often hold structured supplier data
    for tbl in doc.tables:
        for row in tbl.rows:
            paragraphs.append(" | ".join(c.text.strip() for c in row.cells))
    text = "\n".join(paragraphs).strip()
    meta = {
        "source_file": path.name,
        "mime": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "loader": "python-docx",
        "paragraphs": len(paragraphs),
    }
    return [(text, meta)] if text else []


# ---------- XLSX / XLS ----------

def load_xlsx(path: Path) -> list[tuple[str, dict[str, Any]]]:
    from openpyxl import load_workbook

    wb = load_workbook(str(path), data_only=True, read_only=True)
    out: list[tuple[str, dict[str, Any]]] = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows = []
        for row in ws.iter_rows(values_only=True):
            cleaned = [str(c) for c in row if c is not None]
            if cleaned:
                rows.append(" | ".join(cleaned))
        text = "\n".join(rows).strip()
        if not text:
            continue
        out.append((
            text,
            {
                "source_file": path.name,
                "mime": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                "loader": "openpyxl",
                "sheet": sheet_name,
            },
        ))
    return out


# ---------- HTML ----------

def load_html(path: Path | None = None, content: str | None = None,
              url: str | None = None) -> list[tuple[str, dict[str, Any]]]:
    import html2text

    if path is not None:
        content = path.read_text(encoding="utf-8", errors="ignore")
        source = path.name
    elif url:
        source = url
    else:
        source = "html"

    converter = html2text.HTML2Text()
    converter.ignore_images = True
    converter.ignore_links = False
    converter.body_width = 0
    text = converter.handle(content or "").strip()

    # Pull title if possible
    title = ""
    try:
        from bs4 import BeautifulSoup
        soup = BeautifulSoup(content or "", "html.parser")
        if soup.title and soup.title.string:
            title = soup.title.string.strip()[:200]
    except Exception:
        pass

    meta = {
        "source_file": source,
        "mime": "text/html",
        "loader": "html2text",
        "title": title,
        "url": url or "",
    }
    return [(text, meta)] if text else []


# ---------- JSON ----------

def load_json(path: Path) -> list[tuple[str, dict[str, Any]]]:
    data = json.loads(path.read_text(encoding="utf-8", errors="ignore"))
    # Pretty-print so embeddings can match key/value patterns
    text = json.dumps(data, indent=2, default=str)
    return [(
        text,
        {
            "source_file": path.name,
            "mime": "application/json",
            "loader": "json",
            "keys": list(data.keys())[:20] if isinstance(data, dict) else [],
        },
    )]


# ---------- EML ----------

def load_eml(path: Path) -> list[tuple[str, dict[str, Any]]]:
    raw = path.read_bytes()
    msg = email.message_from_bytes(raw)

    body_parts: list[str] = []
    if msg.is_multipart():
        for part in msg.walk():
            ctype = part.get_content_type()
            if ctype == "text/plain":
                body_parts.append(part.get_payload(decode=True).decode(errors="ignore"))
            elif ctype == "text/html":
                # convert to text
                html_text = part.get_payload(decode=True).decode(errors="ignore")
                body_parts.append(load_html(content=html_text)[0][0] if load_html(content=html_text) else "")
    else:
        body_parts.append(msg.get_payload(decode=True).decode(errors="ignore"))

    text = "\n\n".join(p for p in body_parts if p).strip()
    meta = {
        "source_file": path.name,
        "mime": "message/rfc822",
        "loader": "email",
        "from": str(msg.get("From", ""))[:200],
        "to":   str(msg.get("To", ""))[:200],
        "subject": str(msg.get("Subject", ""))[:200],
        "date": str(msg.get("Date", ""))[:80],
    }
    return [(text, meta)] if text else []


# ---------- Plain text ----------

def load_text(path: Path) -> list[tuple[str, dict[str, Any]]]:
    text = path.read_text(encoding="utf-8", errors="ignore").strip()
    return [(text, {
        "source_file": path.name,
        "mime": mimetypes.guess_type(path.name)[0] or "text/plain",
        "loader": "text",
    })] if text else []


# ---------- URL ----------

def load_url(url: str, timeout: float = 15.0) -> list[tuple[str, dict[str, Any]]]:
    parsed = urlparse(url)
    if parsed.scheme not in {"http", "https"}:
        raise ValueError(f"unsupported scheme: {parsed.scheme}")
    r = httpx.get(url, timeout=timeout, follow_redirects=True,
                  headers={"User-Agent": "supply-chain-risk-ai/0.3"})
    r.raise_for_status()
    ctype = r.headers.get("content-type", "").lower()
    if "html" in ctype:
        return load_html(content=r.text, url=url)
    # treat as plain text
    return [(
        r.text.strip(),
        {"source_file": url, "mime": ctype, "loader": "url", "status": r.status_code},
    )]


# ---------- dispatcher ----------

EXT_LOADERS = {
    ".pdf":  load_pdf,
    ".docx": load_docx,
    ".xlsx": load_xlsx,
    ".xls":  load_xlsx,
    ".html": lambda p: load_html(path=p),
    ".htm":  lambda p: load_html(path=p),
    ".json": load_json,
    ".eml":  load_eml,
    ".txt":  load_text,
    ".md":   load_text,
}

ALLOWED_EXTS = set(EXT_LOADERS.keys())


def load_any(path: Path) -> list[tuple[str, dict[str, Any]]]:
    ext = path.suffix.lower()
    loader = EXT_LOADERS.get(ext)
    if not loader:
        raise ValueError(f"no loader for extension {ext}")
    return loader(path)


__all__ = [
    "ALLOWED_EXTS", "EXT_LOADERS", "load_any",
    "load_pdf", "load_docx", "load_xlsx", "load_html", "load_json",
    "load_eml", "load_text", "load_url",
]
